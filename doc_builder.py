"""
Project Repository: https://github.com/Udayatk/Text-To-Document-Markdown-Pdf-Word-Excell
"""
import re
import pandas as pd
from fpdf import FPDF
from docx import Document
import datetime
import os
import unicodedata
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

def parse_markdown_table(md_text):
    """
    Parse a Markdown table string and return a list of dicts.
    """
    lines = [line.strip() for line in md_text.strip().splitlines() if line.strip()]
    table_lines = [line for line in lines if line.startswith('|') and line.endswith('|')]
    if len(table_lines) < 2:
        return []
    
    header = [h.strip() for h in table_lines[0].strip('|').split('|')]
    rows = []
    # Skip header and separator rows
    for line in table_lines[2:]:
        cells = [cell.strip() for cell in line.strip('|').split('|')]
        if len(cells) == len(header):
            rows.append(dict(zip(header, cells)))
    return rows

def get_output_path(filename):
    return os.path.join(os.getcwd(), filename)

def save_excel(chat_history):
    filename = f"chat_document_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    out_path = get_output_path(filename)
    df = pd.DataFrame(chat_history)
    df.to_excel(out_path, index=False)
    return out_path

def save_markdown(chat_history):
    filename = f"chat_document_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
    out_path = get_output_path(filename)
    if not chat_history:
        return out_path
        
    header = list(chat_history[0].keys())
    with open(out_path, 'w', encoding='utf-8') as f:
        if len(header) == 1:
            for entry in chat_history:
                msg = str(entry[header[0]]).strip()
                if msg:
                    f.write(msg + '\n\n')
        else:
            f.write('| ' + ' | '.join(header) + ' |\n')
            f.write('| ' + ' | '.join(['---'] * len(header)) + ' |\n')
            for entry in chat_history:
                f.write('| ' + ' | '.join(str(entry.get(k, '')) for k in header) + ' |\n')
    return out_path

def save_pdf(chat_history):
    filename = f"chat_document_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    out_path = get_output_path(filename)
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=10)
        if not chat_history:
            pdf.multi_cell(0, 8, "(No content provided)")
            pdf.output(out_path)
            return out_path
        header = list(chat_history[0].keys())
        # If single-column, show each message as a paragraph (no header)
        if len(header) == 1:
            pdf.set_font("Arial", size=10)
            for entry in chat_history:
                msg = str(entry[header[0]]).strip()
                if msg:
                    pdf.multi_cell(0, 8, msg)
                    pdf.ln(2)
        else:
            col_widths = []
            effective_page_width = pdf.w - 2 * pdf.l_margin
            base_col_width = effective_page_width / len(header)
            for h in header:
                max_len = max([len(str(row.get(h, ''))) for row in chat_history] + [len(h)])
                width = max(base_col_width * 0.5, min(base_col_width * 2, max_len * 2.5))
                col_widths.append(width)
            total_width = sum(col_widths)
            col_widths = [w * effective_page_width / total_width for w in col_widths]
            row_height = pdf.font_size + 4
            # Header
            pdf.set_font("Arial", style="B", size=10)
            for i, h in enumerate(header):
                pdf.cell(col_widths[i], row_height, h, border=1, align='C')
            pdf.ln(row_height)
            # Rows
            pdf.set_font("Arial", size=10)
            for entry in chat_history:
                pdf.set_x(pdf.l_margin)
                for i, h in enumerate(header):
                    pdf.cell(col_widths[i], row_height, str(entry.get(h, '')), border=1, align='L')
                pdf.ln(row_height)
        pdf.output(out_path)
        return out_path
    except Exception as e:
        if "Unsupported font" in str(e) or "character" in str(e):
            raise RuntimeError(f"PDF generation failed due to a font/character issue. Ensure you use a font that supports all characters in your text. Error: {e}")
        raise RuntimeError(f"PDF generation failed: {str(e)}")


def save_word(chat_history):
    filename = f"chat_document_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    out_path = get_output_path(filename)
    doc = Document()

    # Always treat chat_history as a list of dicts. If only 'message' keys, make a single-column table.
    if isinstance(chat_history, list) and len(chat_history) > 0:
        # If all rows have only 'message' key, treat as single-column
        if all(list(row.keys()) == ['message'] for row in chat_history):
            header = ['Message']
            table = doc.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            hdr_cell = table.rows[0].cells[0]
            run = hdr_cell.paragraphs[0].add_run('Message')
            run.bold = True
            hdr_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            for entry in chat_history:
                row_cell = table.add_row().cells[0]
                row_cell.text = str(entry['message']).strip()
                row_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            header = list(chat_history[0].keys())
            table = doc.add_table(rows=1, cols=len(header))
            table.style = 'Table Grid'
            column_widths_map = {
                "First Name": Pt(100),
                "Last Name": Pt(100),
                "Age": Pt(50)
            }
            for i, col_name in enumerate(header):
                col = table.columns[i]
                col.width = column_widths_map.get(col_name, Pt(90)) # Default width
            hdr_cells = table.rows[0].cells
            for i, h in enumerate(header):
                run = hdr_cells[i].paragraphs[0].add_run(h.strip())
                run.bold = True
                hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            for entry in chat_history:
                row_cells = table.add_row().cells
                for i, k in enumerate(header):
                    cell = row_cells[i]
                    cell.text = str(entry.get(k, '')).strip()
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        doc.add_paragraph("No chat history provided.")

    # Add title
    title = doc.add_paragraph()
    title.add_run('Employee Roster').bold = True
    title.paragraph_format.space_after = Pt(12)

    if chat_history and isinstance(chat_history, list):
        header = list(chat_history[0].keys())
        table = doc.add_table(rows=1, cols=len(header))
        table.style = 'Table Grid'

        # --- FIXED: Column widths now match the data from your image ---
        column_widths_map = {
            "First Name": Pt(100),
            "Last Name": Pt(100),
            "Age": Pt(50)
        }

        # Apply specific widths to each column
        for i, col_name in enumerate(header):
            col = table.columns[i]
            col.width = column_widths_map.get(col_name, Pt(90)) # Default width

        # Header Row
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(header):
            run = hdr_cells[i].paragraphs[0].add_run(h.strip())
            run.bold = True
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Data Rows
        for entry in chat_history:
            row_cells = table.add_row().cells
            for i, k in enumerate(header):
                cell = row_cells[i]
                cell.text = str(entry.get(k, '')).strip()
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.save(out_path)
    return out_path

# --- How to call the function with the data from your image ---

# FIXED: This data now matches the content of the table in your image
employee_data_from_image = [
    {"First Name": "Steve", "Last Name": "Bolinger", "Age": "36"},
    {"First Name": "Stephanie", "Last Name": "Amanda", "Age": "34"},
    {"First Name": "Sherman", "Last Name": "Danes", "Age": "38"},
    {"First Name": "Sid", "Last Name": "Park", "Age": "35"},
    {"First Name": "Stella", "Last Name": "Cachet", "Age": "33"}
]

# Call save_word with the corrected data
output_file_path = save_word(employee_data_from_image)
print(f"Word document saved to: {output_file_path}")

# You can also test the other functions
# output_pdf_path = save_pdf(employee_data_from_image)
# print(f"PDF document saved to: {output_pdf_path}")