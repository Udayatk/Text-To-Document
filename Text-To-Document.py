import streamlit as st
import io
import csv
from io import StringIO
import pandas as pd
from docx import Document
from fpdf import FPDF
import datetime
import os
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

st.title("Text to Document Builder")
st.write("""
Send a message to the chatbot and export the conversation as Markdown documentation.
""")

# 1. Select document type
st.subheader("Step 1: Choose document type")
doc_type = st.selectbox("Select document type to export:", ["Markdown", "PDF", "Word", "Excel"])
type_map = {"Markdown": "md", "PDF": "pdf", "Word": "docx", "Excel": "xlsx"}

# Advanced export functions
def get_output_path(filename):
    return os.path.join(os.getcwd(), filename)

def save_excel(chat_history):
    filename = f"chat_document_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    out_path = get_output_path(filename)
    # Detect if input is a markdown table
    if isinstance(chat_history, list) and len(chat_history) > 0 and all(list(row.keys()) == ['message'] for row in chat_history):
        lines = [str(entry['message']).strip() for entry in chat_history if str(entry['message']).strip()]
        table_lines = [line for line in lines if line.startswith('|')]
        if table_lines:
            headers = [h.strip() for h in table_lines[0].strip('|').split('|')]
            data_start = 1
            if len(table_lines) > 1 and set(table_lines[1].replace('|','').strip()) <= set('-: '):
                data_start = 2
            data_rows = [
                [cell.strip() for cell in row.strip('|').split('|')]
                for row in table_lines[data_start:]
            ]
            df = pd.DataFrame(data_rows, columns=headers)
            df.to_excel(out_path, index=False)
            return out_path
    # Default: regular export
    df = pd.DataFrame(chat_history)
    df.to_excel(out_path, index=False)
    return out_path

def save_markdown(chat_history):
    filename = f"chat_document_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
    out_path = get_output_path(filename)
    if not chat_history:
        return out_path
    header = list(chat_history[0].keys())
    with open(out_path, 'w', encoding='utf-8') as f:
        if len(header) == 1:
            for entry in chat_history:
                msg = str(entry[header[0]]).strip()
                if msg:
                    f.write(msg + '\n\n')
        else:
            f.write('| ' + ' | '.join(header) + ' |\n')
            f.write('| ' + ' | '.join(['---'] * len(header)) + ' |\n')
            for entry in chat_history:
                f.write('| ' + ' | '.join(str(entry.get(k, '')) for k in header) + ' |\n')
    return out_path

def save_pdf(chat_history):
    filename = f"chat_document_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    out_path = get_output_path(filename)
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=10)
        if not chat_history:
            pdf.multi_cell(0, 8, "(No content provided)")
            pdf.output(out_path)
            return out_path
        header = list(chat_history[0].keys())
        if len(header) == 1:
            pdf.set_font("Arial", size=10)
            for entry in chat_history:
                msg = str(entry[header[0]]).strip()
                if msg:
                    pdf.multi_cell(0, 8, msg)
                    pdf.ln(2)
        else:
            col_widths = []
            effective_page_width = pdf.w - 2 * pdf.l_margin
            base_col_width = effective_page_width / len(header)
            for h in header:
                max_len = max([len(str(row.get(h, ''))) for row in chat_history] + [len(h)])
                width = max(base_col_width * 0.5, min(base_col_width * 2, max_len * 2.5))
                col_widths.append(width)
            total_width = sum(col_widths)
            col_widths = [w * effective_page_width / total_width for w in col_widths]
            row_height = pdf.font_size + 4
            pdf.set_font("Arial", style="B", size=10)
            for i, h in enumerate(header):
                pdf.cell(col_widths[i], row_height, h, border=1, align='C')
            pdf.ln(row_height)
            pdf.set_font("Arial", size=10)
            for entry in chat_history:
                pdf.set_x(pdf.l_margin)
                for i, h in enumerate(header):
                    pdf.cell(col_widths[i], row_height, str(entry.get(h, '')), border=1, align='L')
                pdf.ln(row_height)
        pdf.output(out_path)
        return out_path
    except Exception as e:
        if "Unsupported font" in str(e) or "character" in str(e):
            raise RuntimeError(f"PDF generation failed due to a font/character issue. Ensure you use a font that supports all characters in your text. Error: {e}")
        raise RuntimeError(f"PDF generation failed: {str(e)}")

def save_word(chat_history):
    filename = f"chat_document_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    out_path = get_output_path(filename)
    doc = Document()
    if isinstance(chat_history, list) and len(chat_history) > 0:
        if all(list(row.keys()) == ['message'] for row in chat_history):
            # If input is a single string, split into lines for markdown table detection
            full_text = str(chat_history[0]['message']).strip()
            lines = [line for line in full_text.splitlines() if line.strip()]
            table_lines = [line for line in lines if line.startswith('|')]
            if table_lines:
                # Parse header
                headers = [h.strip() for h in table_lines[0].strip('|').split('|')]
                # Find separator line (with dashes)
                data_start = 1
                if len(table_lines) > 1 and set(table_lines[1].replace('|','').strip()) <= set('-: '):
                    data_start = 2
                # Parse data rows
                data_rows = [
                    [cell.strip() for cell in row.strip('|').split('|')]
                    for row in table_lines[data_start:]
                ]
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    run = hdr_cells[i].paragraphs[0].add_run(header)
                    run.bold = True
                    hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                for row in data_rows:
                    row_cells = table.add_row().cells
                    for i, val in enumerate(row):
                        cell = row_cells[i]
                        cell.text = val
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                # First line bold and large, rest normal
                msg_lines = [line for line in full_text.splitlines() if line.strip()]
                first = True
                for msg in msg_lines:
                    para = doc.add_paragraph()
                    run = para.add_run(msg)
                    if first:
                        run.bold = True
                        run.font.size = Pt(18)
                        first = False
                    else:
                        run.font.size = Pt(12)
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            header = list(chat_history[0].keys())
            table = doc.add_table(rows=1, cols=len(header))
            table.style = 'Table Grid'
            column_widths_map = {
                "First Name": Pt(100),
                "Last Name": Pt(100),
                "Age": Pt(50)
            }
            for i, col_name in enumerate(header):
                col = table.columns[i]
                col.width = column_widths_map.get(col_name, Pt(90))
            hdr_cells = table.rows[0].cells
            for i, h in enumerate(header):
                run = hdr_cells[i].paragraphs[0].add_run(h.strip())
                run.bold = True
                hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            for entry in chat_history:
                row_cells = table.add_row().cells
                for i, k in enumerate(header):
                    cell = row_cells[i]
                    cell.text = str(entry.get(k, '')).strip()
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        doc.add_paragraph("")
    doc.save(out_path)
    return out_path

# 2. Manual text input
rows = []
if doc_type == "Excel":
    st.subheader("Step 2: Type CSV text to convert to Excel document")
    typed_text = st.text_area("Enter your CSV text here:")
    if typed_text:
        if st.button("Process Typed Text"):
            reader = csv.DictReader(StringIO(typed_text))
            rows = [row for row in reader]
            if not reader.fieldnames or len(reader.fieldnames) < 1:
                st.error("CSV must have a header row (e.g. name,age)")
            elif not rows or all(all(v == '' for v in row.values()) for row in rows):
                st.error("CSV must have at least one data row below the header.")
            else:
                st.success("CSV text processed.")
else:
    st.subheader(f"Step 2: Type text to convert to {doc_type} document")
typed_text = st.text_area(f"Enter your text for {doc_type} here:")
if doc_type == "Word":
    st.caption("Tip: You can paste Markdown table text here and it will be exported as a formatted table in Word.")
    if typed_text:
        if st.button("Process Typed Text"):
            if typed_text.strip():
                rows = [{"message": typed_text.strip()}]
                st.success(f"Text processed for {doc_type}.")
            else:
                st.error(f"Please enter at least one line of text for {doc_type}.")

# 3. File upload
if doc_type == "Excel":
    st.subheader("Step 3: Or upload a CSV file")
    uploaded_file = st.file_uploader("Choose a CSV file to convert:", type=["csv", "txt"])
    if uploaded_file:
        content = uploaded_file.getvalue().decode('utf-8')
        reader = csv.DictReader(StringIO(content))
        rows = [row for row in reader]
        if not reader.fieldnames or len(reader.fieldnames) < 1:
            st.error("CSV file must have a header row (e.g. name,age)")
        elif not rows or all(all(v == '' for v in row.values()) for row in rows):
            st.error("CSV file must have at least one data row below the header.")
        else:
            st.success("File processed.")
else:
    st.subheader(f"Step 3: Or upload a text file for {doc_type}")
    uploaded_file = st.file_uploader(f"Choose a text file to convert to {doc_type}:", type=["txt", "md", "docx"])
    if uploaded_file:
        content = uploaded_file.getvalue().decode('utf-8')
        for line in content.splitlines():
            if line.strip():
                rows.append({"message": line.strip()})
        if not rows:
            st.error(f"Text file must have at least one line for {doc_type}.")
        else:
            st.success(f"File processed for {doc_type}.")

# 4. Download/export button
st.subheader(f"Step 4: Download your {doc_type} document")
if rows:
    file_ext = type_map[doc_type]
    mime_map = {
        "md": "text/markdown",
        "pdf": "application/pdf",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    }
    out_path = None
    if doc_type == "Markdown":
        out_path = save_markdown(rows)
    elif doc_type == "PDF":
        out_path = save_pdf(rows)
    elif doc_type == "Word":
        out_path = save_word(rows)
    elif doc_type == "Excel":
        out_path = save_excel(rows)
    if out_path:
        with open(out_path, "rb") as f:
            data = f.read()
        st.download_button(
            label=f"Download {doc_type}",
            data=data,
            file_name=os.path.basename(out_path),
            mime=mime_map[file_ext]
        )
        # Optionally, clean up the file after download (uncomment if desired)
        # os.remove(out_path)
else:
    st.info("Process some text or upload a file to enable download.")